import json
import os
import re
from io import BytesIO

import streamlit as st
from docx import Document
from google import genai
from pypdf import PdfReader


# =========================
# Configuration
# =========================
APP_TITLE = "AI Resume ATS Analyzer"
MODEL_NAME = "gemini-3.5-flash"
MAX_FILE_SIZE_MB = 10
MAX_RESUME_CHARS = 50000


# =========================
# Page configuration
# =========================
st.set_page_config(
    page_title=APP_TITLE,
    page_icon="📄",
    layout="wide",
)

st.markdown(
    """
    <style>
        .title {
            text-align: center;
            font-size: 2.6rem;
            font-weight: 700;
            margin-bottom: 0.25rem;
        }
        .subtitle {
            text-align: center;
            color: #666;
            margin-bottom: 2rem;
        }
        .score {
            text-align: center;
            font-size: 3.5rem;
            font-weight: 700;
        }
        .muted {
            color: #666;
            font-size: 0.9rem;
        }
    </style>
    """,
    unsafe_allow_html=True,
)


# =========================
# API key
# =========================
def get_api_key() -> str | None:
    """Read Gemini API key from Streamlit secrets or environment."""
    try:
        key = st.secrets.get("GEMINI_API_KEY")
        if key:
            return str(key).strip()
    except Exception:
        pass

    key = os.getenv("GEMINI_API_KEY")
    return key.strip() if key else None


# =========================
# File extraction
# =========================
def extract_pdf_text(file_bytes: bytes) -> str:
    reader = PdfReader(BytesIO(file_bytes))
    pages = []

    for page in reader.pages:
        try:
            text = page.extract_text()
        except Exception:
            text = None
        if text:
            pages.append(text)

    return "\n".join(pages).strip()


def extract_docx_text(file_bytes: bytes) -> str:
    document = Document(BytesIO(file_bytes))
    parts = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    # Include table text because many resumes use tables for layout.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def extract_resume_text(uploaded_file) -> str:
    file_bytes = uploaded_file.getvalue()
    filename = uploaded_file.name.lower()

    if filename.endswith(".pdf"):
        return extract_pdf_text(file_bytes)

    if filename.endswith(".docx"):
        return extract_docx_text(file_bytes)

    raise ValueError("Only PDF and DOCX files are supported.")


# =========================
# Local deterministic checks
# =========================
def local_resume_checks(text: str) -> dict:
    lower = text.lower()

    section_keywords = {
        "Contact Information": ["email", "phone", "linkedin", "github"],
        "Summary / Profile": ["summary", "objective", "profile"],
        "Experience": ["experience", "employment", "work history"],
        "Education": ["education", "university", "college", "academic"],
        "Skills": ["skills", "technical skills", "core competencies"],
        "Projects": ["projects", "project experience"],
        "Certifications": ["certifications", "certificates"],
    }

    sections = {
        section: any(keyword in lower for keyword in keywords)
        for section, keywords in section_keywords.items()
    }

    email_found = bool(
        re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)
    )

    phone_found = bool(
        re.search(r"(?:\+?\d[\d\s().-]{7,}\d)", text)
    )

    words = re.findall(r"\b[\w+#./-]+\b", text)

    return {
        "word_count": len(words),
        "character_count": len(text),
        "has_email": email_found,
        "has_phone": phone_found,
        "sections": sections,
    }


# =========================
# Gemini analysis
# =========================
def analyze_resume(resume_text: str, job_description: str, api_key: str) -> dict:
    client = genai.Client(api_key=api_key)

    prompt = f"""
You are an expert ATS resume evaluator and career coach.

Analyze the resume and optionally compare it with the supplied job description.

Important rules:
- ATS score must be an estimated score from 0 to 100.
- It is NOT an official score from any ATS vendor.
- Focus only on job-relevant content, structure, readability, keywords,
  skills, experience, achievements, and common ATS practices.
- Do not judge protected or personal characteristics.
- Do not invent jobs, skills, degrees, dates, achievements, or facts.
- If a job description is supplied, identify relevant and missing keywords.
- Give practical, specific improvements.
- Keep rewritten examples faithful to the original resume.

JOB DESCRIPTION:
{job_description.strip() if job_description.strip() else "Not provided."}

RESUME:
{resume_text}

Return ONLY valid JSON in exactly this structure:
{{
  "ats_score": 0,
  "score_summary": "Short explanation of the score.",
  "category_scores": {{
    "ats_formatting": 0,
    "keywords": 0,
    "experience": 0,
    "skills": 0,
    "education": 0,
    "achievements": 0,
    "contact_information": 0,
    "readability": 0
  }},
  "strengths": ["..."],
  "critical_improvements": ["..."],
  "keyword_suggestions": ["..."],
  "missing_keywords": ["..."],
  "formatting_issues": ["..."],
  "content_improvements": ["..."],
  "rewritten_examples": [
    {{
      "original": "...",
      "improved": "..."
    }}
  ],
  "missing_sections": ["..."],
  "final_recommendation": "..."
}}
"""

    response = client.models.generate_content(
        model=MODEL_NAME,
        contents=prompt,
        config={"response_mime_type": "application/json"},
    )

    raw = getattr(response, "text", None)
    if not raw:
        raise RuntimeError("Gemini returned an empty response.")

    try:
        result = json.loads(raw)
    except json.JSONDecodeError as exc:
        raise RuntimeError(
            "Gemini returned invalid JSON. Please click Analyze again."
        ) from exc

    if not isinstance(result, dict):
        raise RuntimeError("Gemini returned an unexpected response.")

    return result


# =========================
# Display helpers
# =========================
def safe_score(value, default=0) -> int:
    try:
        score = int(float(value))
    except (TypeError, ValueError):
        score = default
    return max(0, min(100, score))


def display_bullets(items):
    if not items:
        st.write("None identified.")
        return

    for item in items:
        st.write(f"• {item}")


# =========================
# UI
# =========================
st.markdown(
    '<div class="title">📄 AI Resume ATS Analyzer</div>',
    unsafe_allow_html=True,
)
st.markdown(
    '<div class="subtitle">'
    "Upload your resume, optionally add a job description, and get an AI-powered ATS analysis."
    "</div>",
    unsafe_allow_html=True,
)

with st.sidebar:
    st.header("⚙️ Settings")
    st.write(f"AI model: `{MODEL_NAME}`")
    st.write(f"Maximum file size: `{MAX_FILE_SIZE_MB} MB`")
    st.info(
        "Keep your Gemini API key in Streamlit Secrets or an environment variable. "
        "Never commit the key to GitHub."
    )

uploaded_file = st.file_uploader(
    "📤 Upload your resume",
    type=["pdf", "docx"],
    help=f"PDF or DOCX only. Maximum size: {MAX_FILE_SIZE_MB} MB.",
)

job_description = st.text_area(
    "🎯 Optional: Paste the job description",
    height=220,
    placeholder=(
        "Paste the target job description here for job-specific "
        "keyword matching and recommendations."
    ),
)

if uploaded_file is None:
    st.info("Upload a PDF or DOCX resume to begin.")
    st.stop()

file_size_mb = len(uploaded_file.getvalue()) / (1024 * 1024)
if file_size_mb > MAX_FILE_SIZE_MB:
    st.error(
        f"The selected file is {file_size_mb:.1f} MB. "
        f"Please upload a file smaller than {MAX_FILE_SIZE_MB} MB."
    )
    st.stop()

st.success(f"Ready to analyze: **{uploaded_file.name}**")

if st.button("🚀 Analyze Resume", type="primary", use_container_width=True):
    api_key = get_api_key()

    if not api_key:
        st.error(
            "Gemini API key not found. Add GEMINI_API_KEY to Streamlit Secrets "
            "or set it as an environment variable."
        )
        st.stop()

    with st.spinner("Extracting resume text and analyzing it with Gemini..."):
        try:
            resume_text = extract_resume_text(uploaded_file)

            if not resume_text:
                st.error(
                    "No readable text was found. If this is a scanned/image-only PDF, "
                    "use a text-based PDF or DOCX."
                )
                st.stop()

            resume_for_ai = resume_text[:MAX_RESUME_CHARS]
            local_checks = local_resume_checks(resume_for_ai)

            analysis = analyze_resume(
                resume_for_ai,
                job_description,
                api_key,
            )

        except Exception as exc:
            st.error(f"Analysis failed: {exc}")
            st.stop()

    # Score
    st.divider()
    st.header("📊 ATS Score")

    score = safe_score(analysis.get("ats_score"))
    st.markdown(
        f'<div class="score">{score}/100</div>',
        unsafe_allow_html=True,
    )
    st.progress(score / 100)

    st.write(
        analysis.get(
            "score_summary",
            "No score explanation was returned.",
        )
    )
    st.caption(
        "This is an AI estimate based on common ATS practices, not an official ATS-vendor score."
    )

    # Category scores
    st.header("📈 Category Scores")
    category_scores = analysis.get("category_scores", {})
    category_names = [
        ("ATS Formatting", "ats_formatting"),
        ("Keywords", "keywords"),
        ("Experience", "experience"),
        ("Skills", "skills"),
        ("Education", "education"),
        ("Achievements", "achievements"),
        ("Contact", "contact_information"),
        ("Readability", "readability"),
    ]

    cols = st.columns(4)
    for index, (label, key) in enumerate(category_names):
        with cols[index % 4]:
            value = safe_score(category_scores.get(key))
            st.metric(label, f"{value}/100")

    # Local checks
    st.header("🔎 Basic Resume Checks")
    c1, c2, c3 = st.columns(3)

    with c1:
        st.metric("Word Count", local_checks["word_count"])
    with c2:
        st.metric("Email Found", "Yes" if local_checks["has_email"] else "No")
    with c3:
        st.metric("Phone Found", "Yes" if local_checks["has_phone"] else "No")

    st.subheader("Resume Sections")
    for section, found in local_checks["sections"].items():
        st.write(f"{'✅' if found else '⚠️'} {section}")

    # Strengths and improvements
    st.header("💪 Strengths")
    display_bullets(analysis.get("strengths", []))

    st.header("🚨 Critical Improvements")
    display_bullets(analysis.get("critical_improvements", []))

    st.header("🔑 Keyword Suggestions")
    keywords = analysis.get("keyword_suggestions", [])
    if keywords:
        st.write(" ".join(f"`{str(k)}`" for k in keywords))
    else:
        st.write("No keyword suggestions were returned.")

    if job_description.strip():
        st.header("🎯 Missing Job Keywords")
        display_bullets(analysis.get("missing_keywords", []))

    st.header("📝 Formatting Issues")
    display_bullets(analysis.get("formatting_issues", []))

    st.header("✍️ Content Improvements")
    display_bullets(analysis.get("content_improvements", []))

    st.header("✨ Before / After Examples")
    examples = analysis.get("rewritten_examples", [])

    if examples:
        for example in examples:
            if not isinstance(example, dict):
                continue
            with st.container(border=True):
                st.markdown("**Before:**")
                st.write(example.get("original", ""))
                st.markdown("**After:**")
                st.write(example.get("improved", ""))
    else:
        st.write("No rewritten examples were returned.")

    st.header("📂 Missing / Recommended Sections")
    display_bullets(analysis.get("missing_sections", []))

    st.header("🎯 Final Recommendation")
    st.info(
        analysis.get(
            "final_recommendation",
            "No final recommendation was returned.",
        )
    )

    st.divider()
    st.markdown(
        '<div class="muted">'
        "Privacy note: resume text is sent to Gemini for analysis. "
        "Do not upload sensitive information unless you understand the "
        "privacy policies of the services you use."
        "</div>",
        unsafe_allow_html=True,
    )
